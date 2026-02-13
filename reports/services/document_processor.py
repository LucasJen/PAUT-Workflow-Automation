from docx import Document
import os

class WordTemplateProcessor:
    """ 
    Word Template Processor is used as a find and replace for the selected word document.
    It will search for placeholder matches throughout the document, including tables, and 
    replace them with the provided replacement text
    """

    def __init__(self, template_path, output_path):
        self.template_path = template_path
        self.output_path = output_path
        self.document = Document(template_path)

    def replace(self, placeholder, replacement):
    #Called to replaced all placeholders throughout entire docoument
        for paragraph in self.document.paragraphs:
            self.replace_text_in_paragraphs(paragraph, placeholder, replacement)

        self.replace_in_table_cells_split_runs(placeholder, replacement)

    def replace_text_in_paragraphs(self, paragraph, placeholder, replacement):
    # Replaces text in paragraphs
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, replacement)

    def replace_in_table_cells_split_runs(self, placeholder, replacement):
    # Replaces text in tables, accounts for text in split rows, which can happen when cells are merged.
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text = "".join(run.text for run in paragraph.runs)
                        if placeholder in full_text:
                            new_text = full_text.replace(placeholder, replacement)
                            for run in paragraph.runs:
                                run.text = ""
                            paragraph.runs[0].text = new_text

    def save(self):
        os.makedirs(os.path.dirname(self.output_path), exist_ok=True)
        self.document.save(self.output_path)
